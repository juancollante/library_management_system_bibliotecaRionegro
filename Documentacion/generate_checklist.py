#!/usr/bin/env python3
"""
Generador de Lista de Chequeo para Migración de Software
Biblioteca Rionegro - Formato DOCX
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_checkbox(paragraph):
    """Agrega un checkbox (☐) al inicio del párrafo"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = '☐ ' + run.text

def set_cell_border(cell, **kwargs):
    """
    Establece bordes para una celda de tabla
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Crear elemento de bordes
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)

def create_checklist_document():
    """Crea el documento Word con la lista de chequeo"""
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # TÍTULO PRINCIPAL
    title = doc.add_heading('LISTA DE CHEQUEO PARA REVISIÓN DE MIGRACIÓN DE SOFTWARE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Sistema de Gestión Bibliotecaria - Biblioteca Rionegro', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espacio
    
    # ========== SECCIÓN 1 ==========
    doc.add_heading('1. INFORMACIÓN GENERAL DE LA MIGRACIÓN', level=1)
    
    doc.add_heading('1.1 Datos del Proyecto', level=2)
    p = doc.add_paragraph('Nombre del proyecto: ', style='List Bullet')
    p.add_run('Library Management System Biblioteca Rionegro').bold = True
    
    doc.add_paragraph('Tipo de migración: Migración de aplicación monolítica local a arquitectura containerizada en la nube', style='List Bullet')
    doc.add_paragraph('Plataforma origen: Servidor local tradicional (XAMPP/similar)', style='List Bullet')
    doc.add_paragraph('Plataforma destino: Contenedores Docker + Cloud Platform (Fly.io/AWS/Azure)', style='List Bullet')
    
    p = doc.add_paragraph('Stack tecnológico:', style='List Bullet')
    doc.add_paragraph('Backend: Node.js + Express + MongoDB', style='List Bullet 2')
    doc.add_paragraph('Frontend: React (SPA)', style='List Bullet 2')
    doc.add_paragraph('Base de datos: MongoDB (Atlas o similar en cloud)', style='List Bullet 2')
    
    # 1.2 Alcance
    doc.add_heading('1.2 Alcance de la Migración', level=2)
    
    doc.add_heading('Backend (Node.js/Express)', level=3)
    doc.add_paragraph('Componentes a migrar:', style='List Bullet')
    items = [
        'API REST completa (/api/auth, /api/books, /api/transactions, etc.)',
        '9 módulos de rutas (auth, books, categories, notifications, reservations, search, statistics, transactions, users)',
        '4 modelos de datos (User, Book, BookCategory, BookTransaction)',
        'Sistema de autenticación JWT',
        'Upload de archivos con Multer',
        'Conexión MongoDB con Mongoose'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet 2')
    
    doc.add_paragraph('Adaptaciones necesarias:', style='List Bullet')
    adaptaciones_backend = [
        'Migrar de MongoDB local a MongoDB Atlas (cloud)',
        'Actualizar variables de entorno (.env) para producción',
        'Containerizar con Docker (ya existe Dockerfile)',
        'Configurar CORS para nuevo dominio frontend',
        'Ajustar manejo de archivos estáticos/uploads para cloud storage',
        'Implementar health checks para contenedores'
    ]
    for item in adaptaciones_backend:
        doc.add_paragraph(item, style='List Bullet 2')
    
    doc.add_heading('Frontend (React)', level=3)
    doc.add_paragraph('Componentes a migrar:', style='List Bullet')
    items_front = [
        '12 componentes UI (Header, Footer, ImageSlider, News, PhotoGallery, etc.)',
        'Context API para autenticación (AuthContext, AuthReducer, AuthActions)',
        '3 páginas principales (Home, Allbooks, Signin)',
        '2 dashboards (Admin y Member) con sub-componentes',
        'Build estático optimizado'
    ]
    for item in items_front:
        doc.add_paragraph(item, style='List Bullet 2')
    
    doc.add_paragraph('Adaptaciones necesarias:', style='List Bullet')
    adaptaciones_front = [
        'Actualizar REACT_APP_API_URL para apuntar al backend en cloud',
        'Containerizar con Docker multi-stage build (ya existe Dockerfile)',
        'Servir con nginx optimizado para SPA',
        'Configurar rutas para React Router en producción',
        'Optimizar assets (imágenes en public/assets/images/)'
    ]
    for item in adaptaciones_front:
        doc.add_paragraph(item, style='List Bullet 2')
    
    doc.add_heading('Base de Datos', level=3)
    doc.add_paragraph('Migración de datos:', style='List Bullet')
    items_bd = [
        'Colecciones: users, books, bookcategories, booktransactions, reservations',
        'Exportar desde MongoDB local con mongodump',
        'Importar a MongoDB Atlas con mongorestore',
        'Verificar índices y relaciones'
    ]
    for item in items_bd:
        doc.add_paragraph(item, style='List Bullet 2')
    
    # 1.3 Objetivos
    doc.add_heading('1.3 Objetivos de la Migración', level=2)
    objetivos = [
        'Escalabilidad: Permitir auto-scaling basado en demanda',
        'Disponibilidad: 99.9% uptime con redundancia',
        'Rendimiento: Mejorar tiempos de respuesta aprovechando CDN y caching',
        'Seguridad: HTTPS obligatorio, secrets management, backups automáticos',
        'Mantenimiento: CI/CD automatizado para deploys sin downtime',
        'Costos: Optimizar recursos cloud (contenedores eficientes)'
    ]
    for obj in objetivos:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(obj)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 2 ==========
    doc.add_heading('2. REVISIÓN DE COMPONENTES BACKEND', level=1)
    
    doc.add_heading('2.1 Configuración y Entorno', level=2)
    items_config = [
        '.env migrado: Variables adaptadas para producción',
        'MONGO_URL: Conexión a MongoDB Atlas',
        'PORT: Adaptado a cloud platform (process.env.PORT || 5000)',
        'JWT_SECRET_KEY: Generado nuevo secret para producción (mínimo 32 caracteres)',
        'NODE_ENV=production',
        '.env.example: Actualizado con nuevas variables cloud',
        'Secrets management: Variables sensibles en plataforma cloud (no en repo)'
    ]
    for item in items_config:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('2.2 Dependencias (package.json)', level=2)
    items_deps = [
        'Versiones actualizadas: express, mongoose, bcrypt, cors, multer, jsonwebtoken',
        'Auditoría: npm audit fix ejecutado, sin vulnerabilidades críticas',
        'Lock file: package-lock.json actualizado'
    ]
    for item in items_deps:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('2.3 Código del Servidor (server.js)', level=2)
    items_server = [
        'Puerto dinámico: const PORT = process.env.PORT || 5000;',
        'Conexión MongoDB: Connection string desde variable de entorno',
        'Retry logic implementado (reintentos si falla conexión)',
        'Graceful shutdown (cerrar conexiones al recibir SIGTERM)',
        'CORS configurado para dominio específico',
        'Logs estructurados: Winston/Pino en lugar de console.log',
        'Health check endpoint: GET /health retorna status de BD y API'
    ]
    for item in items_server:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('2.4 Modelos de Datos', level=2)
    items_models = [
        'User.js: Índices en email (unique), validaciones funcionando, hooks bcrypt activos',
        'Book.js: Índices en campos de búsqueda (title, author, isbn)',
        'BookCategory.js: Categorías migradas completas',
        'BookTransaction.js: Referencias a User y Book válidas'
    ]
    for item in items_models:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('2.5 Rutas y APIs', level=2)
    items_routes = [
        '/api/auth: Login con JWT, refresh tokens, rate limiting',
        '/api/books: CRUD completo, paginación, upload de imágenes',
        '/api/categories: Listado y creación de categorías',
        '/api/transactions: Préstamos, devoluciones, validación disponibilidad',
        '/api/reservations: Sistema de reservas activo',
        '/api/users: Gestión de perfiles, roles admin/member',
        '/api/search: Búsqueda fulltext optimizada',
        '/api/statistics: Reportes de préstamos y usuarios',
        '/api/notifications: Notificaciones por email/push'
    ]
    for item in items_routes:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('2.6 Dockerización Backend', level=2)
    items_docker_be = [
        'Dockerfile optimizado: Multi-stage build, imagen node:18-alpine',
        'NODE_ENV=production',
        'npm ci --only=production (sin devDependencies)',
        'Usuario no-root para seguridad',
        'HEALTHCHECK definido',
        '.dockerignore: Excluye node_modules, .env, .git, *.log',
        'Construcción exitosa: docker build -t biblioteca-backend:latest .',
        'Prueba local: docker run con variables de entorno'
    ]
    for item in items_docker_be:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 3 ==========
    doc.add_heading('3. REVISIÓN DE COMPONENTES FRONTEND', level=1)
    
    doc.add_heading('3.1 Configuración', level=2)
    items_config_fe = [
        '.env actualizado: REACT_APP_API_URL apunta a backend cloud',
        'Build de producción: npm run build genera bundle optimizado',
        'Code splitting habilitado',
        'Assets minificados'
    ]
    for item in items_config_fe:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('3.2 Dependencias', level=2)
    items_deps_fe = [
        'React: Versión estable (^18.x)',
        'React Router: Compatible con despliegue cloud',
        'Axios/Fetch: Configurado con baseURL desde .env',
        'Sin vulnerabilidades: npm audit limpio'
    ]
    for item in items_deps_fe:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('3.3 Componentes Principales', level=2)
    items_comp = [
        'Header.js: Links a nuevas URLs',
        'Footer.js: Información de contacto actualizada',
        'ImageSlider.js: Imágenes cargando desde CDN o /assets/',
        'PopularBooks.js: API call a /api/books/popular funcional',
        'RecentAddedBooks.js: Endpoint /api/books/recent respondiendo',
        'ReservedBooks.js: Datos de reservas actualizados',
        'Stats.js: Gráficos mostrando estadísticas correctas'
    ]
    for item in items_comp:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('3.4 Autenticación (Context)', level=2)
    items_auth = [
        'AuthContext.js: Estado global persistente (localStorage/sessionStorage)',
        'AuthReducer.js: Acciones LOGIN, LOGOUT, UPDATE_USER funcionando',
        'AuthActions.js: Llamadas a /api/auth/login y /api/auth/register',
        'Manejo de errores 401 (redirect a login)',
        'Interceptor para agregar token en headers'
    ]
    for item in items_auth:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('3.5 Páginas', level=2)
    items_pages = [
        'Home.js: Renderiza componentes correctamente',
        'Allbooks.js: Listado paginado, búsqueda y filtros operativos',
        'Signin.js: Login funcional, validación, redirect a dashboard'
    ]
    for item in items_pages:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('3.6 Dashboards', level=2)
    items_dash = [
        'AdminDashboard.js: Sub-componentes cargando, permisos validados',
        'MemberDashboard.js: Vista de préstamos, historial, renovaciones'
    ]
    for item in items_dash:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('3.7 Dockerización Frontend', level=2)
    items_docker_fe = [
        'Dockerfile multi-stage: Stage 1 build, Stage 2 nginx:alpine',
        'nginx.conf personalizado para React Router (try_files)',
        'Construcción: docker build -t biblioteca-frontend:latest .',
        'Prueba: docker run -p 80:80'
    ]
    for item in items_docker_fe:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 4 ==========
    doc.add_heading('4. CONTENEDORIZACIÓN Y ORQUESTACIÓN', level=1)
    
    doc.add_heading('4.1 Docker Compose (desarrollo local)', level=2)
    items_compose = [
        'Comando docker-compose up -d levanta stack completo',
        'Servicios se comunican correctamente',
        'Volúmenes persisten datos MongoDB'
    ]
    for item in items_compose:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('4.2 Producción (Fly.io / Cloud Platform)', level=2)
    items_prod = [
        'fly.toml (backend) configurado con región y health checks',
        'Secrets configurados: fly secrets set MONGO_URL, JWT_SECRET_KEY',
        'Deploy backend: fly deploy desde /backend',
        'Deploy frontend: Configurado en servicio estático o nginx'
    ]
    for item in items_prod:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    # ========== SECCIÓN 5 ==========
    doc.add_heading('5. BASE DE DATOS', level=1)
    
    doc.add_heading('5.1 Migración de Datos', level=2)
    items_migracion = [
        'Backup origen: mongodump ejecutado',
        'Importar a Atlas: mongorestore completado',
        'Verificación: Conteo de documentos coincide',
        'Relaciones intactas (populate tests)',
        'Índices recreados automáticamente'
    ]
    for item in items_migracion:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('5.2 MongoDB Atlas Configuración', level=2)
    items_atlas = [
        'Cluster creado: M0/M10 según necesidad',
        'IP Whitelist: Configurado para cloud platform',
        'Database User: Creado con permisos readWrite',
        'Backups automáticos: Habilitados',
        'Connection string: Almacenado como secret'
    ]
    for item in items_atlas:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('5.3 Rendimiento', level=2)
    items_perf = [
        'Índices compuestos para búsquedas frecuentes',
        'Connection pooling: maxPoolSize configurado',
        'Queries optimizados con proyecciones',
        'Tiempos de respuesta aceptables (<200ms CRUD básico)'
    ]
    for item in items_perf:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 6 ==========
    doc.add_heading('6. SEGURIDAD', level=1)
    
    doc.add_heading('6.1 Autenticación y Autorización', level=2)
    items_auth_sec = [
        'JWT robusto: Secret 32+ caracteres, expiración configurada',
        'Bcrypt salt rounds: Mínimo 10',
        'Middleware de autorización para roles admin/member',
        'Rate limiting: Configurado en /api/auth'
    ]
    for item in items_auth_sec:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('6.2 Comunicación', level=2)
    items_comm = [
        'HTTPS obligatorio: Certificado SSL/TLS válido',
        'HSTS header configurado',
        'CORS restrictivo a dominios permitidos',
        'Helmet.js: Headers de seguridad activos'
    ]
    for item in items_comm:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('6.3 Datos Sensibles', level=2)
    items_sensitive = [
        'Environment variables: Nunca en código',
        '.env en .gitignore',
        'Secrets rotation: JWT_SECRET rotado post-migración',
        'Logs sanitizados: No loggear passwords, tokens'
    ]
    for item in items_sensitive:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    # ========== SECCIÓN 7 ==========
    doc.add_heading('7. PRUEBAS Y VALIDACIÓN', level=1)
    
    doc.add_heading('7.1 Pruebas Funcionales Post-Migración', level=2)
    items_tests = [
        'Autenticación: Login válido → JWT, Login inválido → 401',
        'CRUD Libros: Admin crea, usuario ve, permisos validados',
        'Transacciones: Préstamo exitoso, devolución funcional',
        'Reservas: Usuario puede reservar, notificaciones activas',
        'Búsqueda: Búsqueda por título, filtros por categoría',
        'Estadísticas: Dashboard admin muestra métricas correctas'
    ]
    for item in items_tests:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('7.2 Pruebas de Integración', level=2)
    items_int = [
        'Flujo completo usuario: Registro → Préstamo → Devolución',
        'Flujo admin: Agregar libro → Ver transacciones → Reportes'
    ]
    for item in items_int:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('7.3 Pruebas de Rendimiento', level=2)
    items_load = [
        'Load testing: Tiempo respuesta promedio < 200ms',
        'Sin errores bajo 100 usuarios concurrentes',
        'Database performance: Queries < 50ms',
        'Connection pool no saturado'
    ]
    for item in items_load:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('7.4 Pruebas Automatizadas (Postman/Newman)', level=2)
    items_newman = [
        'Colección Postman actualizada con URLs cloud',
        'Tests de assertions (status, schema, tiempos)',
        'Ejecución en CI/CD: newman run completado exitosamente'
    ]
    for item in items_newman:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 8 ==========
    doc.add_heading('8. MONITOREO Y LOGS', level=1)
    
    doc.add_heading('8.1 Logging', level=2)
    items_log = [
        'Winston configurado con formato JSON',
        'Logs en cloud: Integrados con servicio (CloudWatch/Datadog)',
        'Rotación automática: Logs antiguos archivados'
    ]
    for item in items_log:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('8.2 Monitoreo', level=2)
    items_monitor = [
        'Uptime monitoring: UptimeRobot/Pingdom configurado',
        'Alertas: Email si error rate > 5%',
        'APM (opcional): New Relic/Datadog para métricas'
    ]
    for item in items_monitor:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    # ========== SECCIÓN 9 ==========
    doc.add_heading('9. DESPLIEGUE', level=1)
    
    doc.add_heading('9.1 Plataforma de Hosting', level=2)
    doc.add_paragraph('Opción seleccionada: Fly.io (según fly.toml existente)', style='List Bullet')
    items_hosting = [
        'Backend desplegado en Fly.io: fly deploy ejecutado',
        'URL backend: https://biblioteca-backend.fly.dev',
        'Health check OK',
        'Frontend desplegado en Vercel/Netlify o Fly.io'
    ]
    for item in items_hosting:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('9.2 CI/CD Pipeline', level=2)
    items_cicd = [
        'GitHub Actions configurado: Deploy automático en push a main',
        'Tests automáticos antes de deploy (Lint, Unit tests, Newman)',
        'Notificaciones de deploy exitoso/fallido'
    ]
    for item in items_cicd:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('9.3 Rollback Strategy', level=2)
    items_rollback = [
        'Versionado de deploys: Tags en Docker images',
        'Rollback rápido: fly releases rollback disponible',
        'Procedimiento documentado'
    ]
    for item in items_rollback:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('9.4 Dominios y DNS', level=2)
    items_dns = [
        'Dominio registrado: biblioteca-rionegro.com',
        'DNS configurado: A record y CNAME',
        'SSL/TLS: Certificado Let\'s Encrypt auto-renovado'
    ]
    for item in items_dns:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 10 ==========
    doc.add_heading('10. DOCUMENTACIÓN', level=1)
    
    doc.add_heading('10.1 Documentación Técnica', level=2)
    items_docs = [
        'README.md actualizado: Arquitectura cloud, deploy commands',
        'DEPLOYMENT.md: Paso a paso Fly.io y MongoDB Atlas',
        'ARCHITECTURE.md: Diagrama de componentes y flujo de datos'
    ]
    for item in items_docs:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('10.2 API Documentation', level=2)
    items_api_docs = [
        'Swagger/OpenAPI implementado en /api-docs',
        'Postman Collection publicada: Compartida en workspace público'
    ]
    for item in items_api_docs:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    # ========== SECCIÓN 11 ==========
    doc.add_heading('11. CRITERIOS DE ACEPTACIÓN FINAL', level=1)
    
    doc.add_heading('11.1 Funcional', level=2)
    items_accept_func = [
        '100% de funcionalidades originales operativas en cloud',
        'No regresiones: Tests de regresión pasados',
        'Performance igual o mejor que versión local'
    ]
    for item in items_accept_func:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('11.2 No Funcional', level=2)
    items_accept_nf = [
        'Disponibilidad: 99.9% uptime en primeros 30 días',
        'Tiempo de respuesta API: < 300ms p95',
        'Frontend (FCP): < 2s',
        'HTTPS obligatorio',
        'Sin vulnerabilidades críticas',
        'Backups diarios activos'
    ]
    for item in items_accept_nf:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('11.3 Datos', level=2)
    items_accept_data = [
        'Migración completa: Todos los documentos transferidos',
        'Integridad: Queries de validación exitosas',
        'Pérdida de datos: 0%'
    ]
    for item in items_accept_data:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_page_break()
    
    # ========== SECCIÓN 12 ==========
    doc.add_heading('12. PLAN DE CONTINGENCIA', level=1)
    
    doc.add_heading('12.1 Rollback Plan', level=2)
    items_contingencia = [
        'Mantener entorno local activo durante 30 días post-migración',
        'Backup completo pre-migración disponible',
        'Procedimiento documentado: Cómo volver a local en < 2 horas'
    ]
    for item in items_contingencia:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    doc.add_heading('12.2 Soporte Post-Migración', level=2)
    items_soporte = [
        'Monitoreo intensivo primeros 7 días (24/7)',
        'Canal de soporte: Email/Slack para reportar issues',
        'Hotfixes prioritarios: SLA < 4 horas para críticos'
    ]
    for item in items_soporte:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(item)
    
    # ========== SECCIÓN 13 ==========
    doc.add_heading('13. CRONOGRAMA EJECUTIVO', level=1)
    
    # Crear tabla de cronograma
    table = doc.add_table(rows=11, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    headers = ['Fase', 'Actividad', 'Duración', 'Responsable', 'Estado']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Datos
    cronograma_data = [
        ['1', 'Setup MongoDB Atlas + migración datos', '2 días', 'DevOps', 'Pendiente'],
        ['2', 'Actualizar configs backend (.env, CORS)', '1 día', 'Backend Dev', 'Pendiente'],
        ['3', 'Dockerizar backend + pruebas locales', '2 días', 'Backend Dev', 'Pendiente'],
        ['4', 'Deploy backend a Fly.io', '1 día', 'DevOps', 'Pendiente'],
        ['5', 'Actualizar frontend (API URLs)', '1 día', 'Frontend Dev', 'Pendiente'],
        ['6', 'Dockerizar frontend + deploy', '2 días', 'Frontend/DevOps', 'Pendiente'],
        ['7', 'Configurar DNS + SSL', '1 día', 'DevOps', 'Pendiente'],
        ['8', 'Pruebas end-to-end en producción', '2 días', 'QA', 'Pendiente'],
        ['9', 'Monitoreo + ajustes', '3 días', 'Todo el equipo', 'Pendiente'],
        ['10', 'Go-live oficial', '-', 'Product Owner', 'Pendiente']
    ]
    
    for i, row_data in enumerate(cronograma_data, start=1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph('Duración total estimada: 15 días hábiles').bold = True
    
    doc.add_page_break()
    
    # ========== SECCIÓN 14 - FIRMAS ==========
    doc.add_heading('14. FIRMAS Y APROBACIONES', level=1)
    
    # Tabla de firmas
    firma_table = doc.add_table(rows=8, cols=4)
    firma_table.style = 'Light Grid Accent 1'
    
    # Encabezados
    firma_headers = ['Rol', 'Nombre', 'Firma', 'Fecha']
    for i, header in enumerate(firma_headers):
        cell = firma_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Roles
    roles = [
        'Desarrollador Backend',
        'Desarrollador Frontend',
        'DevOps Engineer',
        'QA/Tester',
        'DBA',
        'Líder Técnico',
        'Product Owner'
    ]
    
    for i, rol in enumerate(roles, start=1):
        firma_table.rows[i].cells[0].text = rol
    
    doc.add_paragraph()
    doc.add_heading('Observaciones específicas de esta migración:', level=2)
    observaciones = [
        'Backend ya tiene Dockerfile y fly.toml → migración más rápida',
        'Frontend tiene build/ pre-generado → validar si está actualizado',
        'Considerar uso de CDN (CloudFront/Cloudflare) para assets estáticos',
        'Evaluar implementar Redis para sessions/caching (optimización futura)',
        'Multer uploads → migrar a S3/Azure Blob si volumen crece'
    ]
    for obs in observaciones:
        doc.add_paragraph(obs, style='List Number')
    
    doc.add_paragraph()
    doc.add_heading('Estado de la migración:', level=2)
    estados = [
        'En planificación',
        'En ejecución',
        'Completada exitosamente',
        'Completada con observaciones',
        'Requiere revisión'
    ]
    for estado in estados:
        p = doc.add_paragraph(style='List Bullet')
        add_checkbox(p)
        p.add_run(estado)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer_text = doc.add_paragraph()
    footer_text.add_run('Checklist específico para: ').italic = True
    footer_text.add_run('Library Management System - Biblioteca Rionegro').bold = True
    
    footer_text2 = doc.add_paragraph()
    footer_text2.add_run('Stack: Node.js + Express + MongoDB → Docker + Fly.io/Cloud').italic = True
    
    footer_text3 = doc.add_paragraph()
    footer_text3.add_run('Versión: 2.0 - Migración a Cloud').italic = True
    
    footer_text4 = doc.add_paragraph()
    footer_text4.add_run('Fecha: Diciembre 2025').italic = True
    
    # Guardar documento
    output_path = 'Lista_Chequeo_Migracion_BibliotecaRionegro.docx'
    doc.save(output_path)
    print(f"✓ Documento generado exitosamente: {output_path}")
    return output_path

if __name__ == "__main__":
    try:
        create_checklist_document()
    except Exception as e:
        print(f"Error al generar el documento: {e}")
        import traceback
        traceback.print_exc()
